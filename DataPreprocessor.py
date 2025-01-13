import pandas as pd
import re
import emoji
import pytest
import numpy as np
from docx import Document
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


def datapreprocessor(df):
    # Ensure the 'comment' column exists
    if 'comment' not in df.columns:
        raise ValueError("The DataFrame must contain a column named 'comment'.")
    
    # Remove empty comments (NaN or empty strings)
    df = df[df['comment'].notna()] 
    df['comment'] = df['comment'].apply(lambda x: x.strip() if isinstance(x, str) else '')  
    df = df[df['comment'] != ""]  
    
  
    # Remove URLs (http/https links)
    df['comment'] = df['comment'].apply(lambda x: re.sub(r"http\S+|https\S+", "", x)) 
    
    # Normalize emojis (convert to words)
    def replace_emojis(text):
        return emoji.demojize(text, delimiters=("", "")) if isinstance(text, str) else text
    df['comment'] = df['comment'].apply(replace_emojis)  
    
    # Convert to lowercase
    df['comment'] = df['comment'].str.lower()
    
    # Remove extra spaces within the comments 
    df['comment'] = df['comment'].apply(lambda x: re.sub(r'\s+', ' ', x) if isinstance(x, str) else x)
    
    # Normalize non-alphanumeric characters 
    df['comment'] = df['comment'].apply(lambda x: re.sub(r'[^\w\s]', '', x) if isinstance(x, str) else x)
    
    # Remove duplicate comments 
    df = df.drop_duplicates(subset='comment', keep='first')  

    df = df.dropna(subset=['comment'])  
    df = df[df['comment'].str.strip().ne('')]  # Remove empty comments
    
    # Reset index for a clean DataFrame
    df = df.reset_index(drop=True)
    
    return df

def test_datapreprocessor(df):
    errors = []  
    
    # Test 1: Check if the 'comment' column exists
    if 'comment' not in df.columns:
        errors.append("The DataFrame must contain a column named 'comment'.")

    # Check for NaN comments
    nan_comments = df['comment'].isna().any()
    if nan_comments:
        errors.append("The DataFrame contains NaN comments.")

    # Check for empty comments
    empty_comments = df['comment'].str.strip().eq('').any()
    if empty_comments:
        errors.append("The DataFrame contains empty comments.")
    
    # Check for duplicates
    duplicates = df['comment'].duplicated().any()
    if duplicates:
        errors.append("The DataFrame contains duplicate comments.")
    
    # Check for URLs
    urls_present = df['comment'].str.contains(r'http[s]?://\S+', regex=True).any()
    if urls_present:
        errors.append("The DataFrame contains URLs.")
    
    # If there are any errors, raise them at once
    if errors:
        raise AssertionError("\n".join(errors))
    
    # If no errors, print that the DataFrame is clean
    print(errors)

def inspect_data(df):
    df = df.rename(columns={'PostContent': 'comment', 'Username': 'username', 'Timestamp': 'time'})

    doc = Document()
    doc.add_heading('Data Overview', 0)
    
    def set_table_borders(table):
        tbl = table._element
        tbl_properties = tbl.tblPr
        tbl_properties.append(parse_xml(r'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="single" w:sz="4"/><w:left w:val="single" w:sz="4"/><w:bottom w:val="single" w:sz="4"/><w:right w:val="single" w:sz="4"/><w:insideH w:val="single" w:sz="4"/><w:insideV w:val="single" w:sz="4"/></w:tblBorders>'))

    doc.add_heading('General Information', level=1)
    doc.add_paragraph(f"Shape of the dataset: {df.shape}")
    doc.add_paragraph(f"Columns: {list(df.columns)}")
    
    doc.add_heading('Data Types', level=2)
    doc.add_paragraph(str(df.dtypes))

    doc.add_heading('Descriptive Statistics (Numerical Columns)', level=2)
    stats = df.describe().transpose()
    table = doc.add_table(rows=1, cols=len(stats.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(stats.columns):
        hdr_cells[i].text = col
    
    for index, row in stats.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
    
    set_table_borders(table)

    doc.add_heading('Most Common Entries in "username"', level=1)
    if 'username' in df.columns:
        doc.add_paragraph(str(df['username'].value_counts().head(10)))

    doc.add_heading('Likes and Dislikes', level=1)
    if 'likes' in df.columns and 'dislikes' in df.columns:
        doc.add_paragraph(f"Total Likes: {df['likes'].sum()}")
        doc.add_paragraph(f"Total Dislikes: {df['dislikes'].sum()}")
        most_liked = df.loc[df['likes'].idxmax()] if 'likes' in df.columns else None
        if most_liked is not None:
            doc.add_paragraph(f"Most Liked Comment: {most_liked['likes']} likes")

    doc.add_heading('Comments Analysis', level=1)
    if 'comment' in df.columns:
        comment_lengths = df['comment'].dropna().apply(lambda x: len(x.split()))
        doc.add_paragraph(f"Average comment length: {np.mean(comment_lengths):.2f} words")
        doc.add_paragraph(f"Longest comment length: {np.max(comment_lengths)} words")

    doc.add_heading('Additional Insights', level=1)
    if 'time' in df.columns:
        df['time'] = pd.to_datetime(df['time'], errors='coerce')
        doc.add_paragraph(f"Time range: {df['time'].min()} to {df['time'].max()}")

    doc.add_heading('Example Rows', level=1)
    example_rows = df.head(2)
    table = doc.add_table(rows=1, cols=len(example_rows.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(example_rows.columns):
        hdr_cells[i].text = col
    
    for index, row in example_rows.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
    
    set_table_borders(table)

    doc.add_heading('Comments Frequency Distribution by Month', level=1)
    if 'time' in df.columns:
        df['month'] = df['time'].dt.to_period('M')
        monthly_counts = df.groupby('month').size()
        frequency_dist = monthly_counts.value_counts().sort_index()
        doc.add_paragraph(f"Frequency Distribution of Comments Per Month: {dict(frequency_dist)}")

    doc.add_heading('Comments Count by Month', level=1)
    if 'time' in df.columns:
        df['month'] = df['time'].dt.to_period('M')
        comments_by_month = df.groupby('month').size()
        for month, count in comments_by_month.items():
            doc.add_paragraph(f"{month}: {count} comments")

    doc.save('data_overview_report.docx')

    print("\n--- Data Overview ---")
    print(f"Shape of the dataset: {df.shape}")
    print(f"Columns: {list(df.columns)}")

    print("\n--- Data Types ---")
    print(df.dtypes)

    print("\n--- Descriptive Statistics (Numerical Columns) ---")
    print(df.describe().transpose())

    print("\n--- Most Common Entries in 'username' ---")
    if 'username' in df.columns:
        print(df['username'].value_counts().head(10))

    print("\n--- Likes and Dislikes ---")
    if 'likes' in df.columns and 'dislikes' in df.columns:
        print(f"Total Likes: {df['likes'].sum()}")
        print(f"Total Dislikes: {df['dislikes'].sum()}")
        most_liked = df.loc[df['likes'].idxmax()] if 'likes' in df.columns else None
        if most_liked is not None:
            print(f"Most Liked Comment: {most_liked['likes']} likes")

    print("\n--- Comments Analysis ---")
    if 'comment' in df.columns:
        comment_lengths = df['comment'].dropna().apply(lambda x: len(x.split()))
        print(f"Average comment length: {np.mean(comment_lengths):.2f} words")
        print(f"Longest comment length: {np.max(comment_lengths)} words")

    print("\n--- Additional Insights ---")
    if 'time' in df.columns:
        df['time'] = pd.to_datetime(df['time'], errors='coerce')
        print(f"Time range: {df['time'].min()} to {df['time'].max()}")

    print("\n--- Example Rows ---")
    print(df.head(2))

    print("\n--- Frequency Distribution of Comments by Month ---")
    if 'time' in df.columns:
        df['month'] = df['time'].dt.to_period('M')
        monthly_counts = df.groupby('month').size()
        frequency_dist = monthly_counts.value_counts().sort_index()
        print(f"Frequency Distribution of Comments Per Month: {dict(frequency_dist)}")

    print("\n--- Comments Count by Month ---")
    if 'time' in df.columns:
        df['month'] = df['time'].dt.to_period('M')
        comments_by_month = df.groupby('month').size()
        for month, count in comments_by_month.items():
            print(f"{month}: {count} comments")



